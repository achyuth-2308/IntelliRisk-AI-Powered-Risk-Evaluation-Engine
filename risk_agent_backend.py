import os
from typing import List
import pandas as pd
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from dotenv import load_dotenv
from docx import Document
from datetime import datetime

# ------------- Configuration: file paths ---------------
docx_path = r"C:\Users\achyuth.m\Documents\knorr_bremse\1_products_spec.docx"
csv_path = r"C:\Users\achyuth.m\Documents\knorr_bremse\2_knorr_bremse_component_reviews.csv"
docx_store_dir = r"C:\Users\achyuth.m\Documents\knorr_bremse\faiss_docx_store"
csv_store_dir = r"C:\Users\achyuth.m\Documents\knorr_bremse\faiss_csv_store"
# -------------------------------------------------------

# ------------- Load environment and models -------------
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

llm = GoogleGenerativeAI(model="gemini-2.5-pro", google_api_key=GOOGLE_API_KEY)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

faiss_store_docx = None
faiss_store_csv = None

# ------------- Helpers for loading documents -------------
def docx_to_documents(docx_path: str) -> List[str]:
    doc = Document(docx_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def csv_to_documents(csv_path: str) -> List[str]:
    df = pd.read_csv(csv_path)
    return [
        "; ".join(f"{col}: {row[col]}" for col in df.columns)
        for _, row in df.iterrows()
    ]

# ------------ Building FAISS stores ---------------------
def add_documents_to_faiss(documents: List[str], which: str) -> None:
    global faiss_store_docx, faiss_store_csv
    if not documents:
        print(f"No documents provided for {which} FAISS DB.")
        return
    try:
        if which == "docx":
            print(f"Adding {len(documents)} documents to DOCX FAISS store...")
            faiss_store_docx = FAISS.from_texts(texts=documents, embedding=embeddings)
        elif which == "csv":
            print(f"Adding {len(documents)} documents to CSV FAISS store...")
            faiss_store_csv = FAISS.from_texts(texts=documents, embedding=embeddings)
        else:
            raise ValueError("Unknown FAISS vector store identifier!")
        print(f"Documents successfully added to {which} FAISS DB.")
    except Exception as e:
        raise Exception(f"Failed to add documents to FAISS ({which}): {e}")

def save_faiss_stores():
    if faiss_store_docx is not None:
        faiss_store_docx.save_local(docx_store_dir)
        print(f"SAVED: DOCX FAISS store at {docx_store_dir}")
    if faiss_store_csv is not None:
        faiss_store_csv.save_local(csv_store_dir)
        print(f"SAVED: CSV FAISS store at {csv_store_dir}")

# ------------ Build or Load FAISS stores (NEW) ----------
def build_or_load_faiss(documents: List[str], store_dir: str, which: str):
    """
    Loads the FAISS store from disk if it exists, otherwise initializes and saves it.
    Returns the FAISS index object.
    """
    if os.path.isdir(store_dir) and os.listdir(store_dir):
        print(f"Loading existing FAISS store for {which} from {store_dir}")
        return FAISS.load_local(store_dir, embeddings, allow_dangerous_deserialization=True)
    else:
        print(f"Building new FAISS store for {which} at {store_dir}")
        index = FAISS.from_texts(documents, embeddings)
        index.save_local(store_dir)
        return index

# ------------ MarkDown Cleaner and Converter -----------
import re

def clean_and_format_markdown(text: str) -> str:
    # Remove all * and # and redundant whitespace
    text = text.replace("*", "")
    lines = text.splitlines()
    clean_lines = []
    for line in lines:
        clean_line = line.replace("#", "").strip()
        clean_lines.append(clean_line)
    md = "\n".join(clean_lines)

    # Replace section names with Markdown headings if needed
    section_names = [
        "Executive Summary",
        "System Specifications / Requirements",
        "Component History & Reviews",
        "Risk Evaluation",
        "Missing Information / Requirements Gap",
    ]
    for sec in section_names:
        # Replace section names with markdown headings if line starts with that section
        md = re.sub(rf"^{sec}\s*:?$", f"## {sec}", md, flags=re.MULTILINE)

    # Risk Score line: move to percentage and bold
    def risk_score_to_pct(line):
        match = re.match(r"Risk Score:\s*([0-9]*\.?[0-9]+)\s*\((Low|Medium|High)\)", line, flags=re.I)
        if match:
            dec = float(match.group(1))
            pct = int(round(dec * 100))
            level = match.group(2)
            return f"**Risk Score: {pct}% ({level})**"
        return line

    md_lines = []
    found_score = False
    for line in md.splitlines():
        if not found_score and line.lower().startswith("risk score"):
            md_lines.append(risk_score_to_pct(line))
            found_score = True
        else:
            md_lines.append(line)
    return "\n".join(md_lines)

# ------------ Risk Evaluator Agent ----------------------
def risk_evaluator(product_name: str, faiss_store_docx, faiss_store_csv, llm, output_report_path: str, return_markdown=False):
    query = str(product_name)
    # Get relevant context from both stores
    spec_docs = faiss_store_docx.similarity_search(query, k=3)
    hist_docs = faiss_store_csv.similarity_search(query, k=3)

    spec_content = "\n".join([doc.page_content for doc in spec_docs]) or "Not found"
    hist_content = "\n".join([doc.page_content for doc in hist_docs]) or "Not found"

    print("\n[Agent] Retrieved specifications:")
    print(spec_content)
    print("\n[Agent] Retrieved component history/reviews:")
    print(hist_content)

    # Prompt LLM for risk evaluation
    prompt = f'''
You are an expert manufacturing product risk evaluator. Analyze the following product/component based on its specifications and operational history.

Component/Product Name:
{product_name}

Specifications/Requirements (from system requirements):
{spec_content}

Component History & User Reviews (from operational, breakage, maintenance, or user review data):
{hist_content}

Produce a detailed risk evaluation report in clearly marked sections:
Executive Summary, System Specifications / Requirements, Component History & Reviews, Risk Evaluation, Missing Information / Requirements Gap.

For Risk Evaluation, start with "Risk Score: <decimal between 0 and 1> (Low/Medium/High)" as the first line. No hashtags, no asterisks, and no bullet points anywhere in your output.
    '''
    report = llm.invoke(prompt)
    try:
        report_text = report if isinstance(report, str) else getattr(report, "text", str(report))
    except Exception:
        report_text = str(report)

    print("\n[Agent] Risk report generated!")

    # --- Clean and convert to markdown ---
    md_report_text = clean_and_format_markdown(report_text)

    # --- Write report as markdown in docx (simulating pretty, organized doc) ---
    doc = Document()
    doc.add_heading(f'Risk Evaluation Report: "{product_name}"', 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for mdline in md_report_text.splitlines():
        # Headings ("## "): use docx headings, Risk Score: bold
        if mdline.startswith("## "):
            doc.add_heading(mdline.lstrip("# ").strip(), level=1)
        elif mdline.startswith("**") and mdline.endswith("**"):
            doc.add_paragraph(mdline.strip("*")).bold = True
        else:
            doc.add_paragraph(mdline)
    doc.save(output_report_path)
    print(f'[Agent] Report saved at {output_report_path}')
    if return_markdown:
        return md_report_text  # For live display in Streamlit etc.
    else:
        return output_report_path

# ------------ Main execution (for CLI) ------------------
if __name__ == "__main__":
    print("Loading documents from:")
    print(" - DOCX:", docx_path)
    print(" - CSV :", csv_path)
    docs_from_docx = docx_to_documents(docx_path)
    docs_from_csv = csv_to_documents(csv_path)
    print(f"DOCX paragraphs: {len(docs_from_docx)}")
    print(f"CSV rows: {len(docs_from_csv)}")

    add_documents_to_faiss(docs_from_docx, which="docx")
    add_documents_to_faiss(docs_from_csv, which="csv")
    save_faiss_stores()

    # --- Run agent for a component ---
    component_or_product = "brake cable"
    report_path = rf"C:\Users\achyuth.m\Documents\knorr_bremse\risk_report_{component_or_product.replace(' ','_')}.docx"
    risk_evaluator(
        product_name=component_or_product,
        faiss_store_docx=faiss_store_docx,
        faiss_store_csv=faiss_store_csv,
        llm=llm,
        output_report_path=report_path,
    )
